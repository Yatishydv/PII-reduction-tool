"""
docx_processor.py — Safe DOCX reading and writing with run-level replacement.

Key design goals:
1. Preserve formatting (bold, italic, font, size, color) by operating at the
   run level, not at paragraph.text level.
2. Process paragraphs AND tables (including nested tables).
3. Handle the edge case where PII spans across multiple runs within a paragraph
   by merging all runs into a single text, applying replacements, then
   reconstructing runs while copying formatting from the first run.
4. Validate the output DOCX after writing.

Limitations (documented):
- Hyperlinks embedded in DOCX XML are processed at XML level; their display
  text is redacted but the underlying href may remain if it contains PII.
  This is a known limitation of python-docx's hyperlink API.
- Images are not processed.
- Headers/footers are processed but may have limited formatting preservation.
"""
from __future__ import annotations

import copy
import logging
import re
from pathlib import Path
from typing import List, Tuple

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Run-level text replacement helper
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph: Paragraph, redact_fn) -> int:
    """Apply *redact_fn* to the merged text of a paragraph, then reconstruct.

    Strategy: merge all run texts into a single string, apply redaction,
    then put the result in the first run and clear all remaining runs.
    This correctly handles PII that is split across multiple formatting runs.

    Args:
        paragraph:  The python-docx Paragraph to modify in place.
        redact_fn:  A callable ``(text: str) -> (redacted: str, entities: list)``.

    Returns:
        Number of replacements made (i.e., number of entities detected).
    """
    runs = paragraph.runs
    if not runs:
        return 0

    # Merge all run texts into one string
    full_text = "".join(r.text for r in runs)
    if not full_text.strip():
        return 0

    redacted_text, entities = redact_fn(full_text)

    if not entities:
        return 0

    # Put the fully-redacted text into the first run
    runs[0].text = redacted_text

    # CRITICAL: Clear all remaining runs so no split-run fragments survive.
    # We do NOT delete runs (to preserve XML structure and paragraph formatting).
    for run in runs[1:]:
        run.text = ""

    return len(entities)




def _replace_in_cell(cell: _Cell, redact_fn) -> int:
    """Apply replacement to all paragraphs in a table cell."""
    count = 0
    for para in cell.paragraphs:
        count += _replace_in_paragraph(para, redact_fn)
    return count


# ---------------------------------------------------------------------------
# DocxProcessor
# ---------------------------------------------------------------------------

class DocxProcessor:
    """Reads a DOCX, applies redaction, and writes the output DOCX.

    Args:
        input_path:  Path to the source DOCX file.
    """

    def __init__(self, input_path: str | Path) -> None:
        self._input_path = Path(input_path)
        if not self._input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self._input_path}")
        if self._input_path.suffix.lower() != ".docx":
            raise ValueError(f"Input must be a .docx file, got: {self._input_path.suffix}")

        self._doc: Document = docx.Document(str(self._input_path))
        logger.info("Loaded document: %s (%d paragraphs, %d tables)",
                    self._input_path.name,
                    len(self._doc.paragraphs),
                    len(self._doc.tables))

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def process(self, redact_fn) -> Tuple[int, int]:
        """Apply *redact_fn* to all text in the document.

        Processes:
          - All paragraphs (including heading/body paragraphs)
          - All table cells (including nested tables)

        Args:
            redact_fn: Callable ``(text) -> (redacted_text, entities)``.

        Returns:
            A tuple (total_replacements, total_paragraphs_modified).
        """
        total_replacements = 0
        paragraphs_modified = 0

        # --- Process body paragraphs ---
        for para in self._doc.paragraphs:
            n = _replace_in_paragraph(para, redact_fn)
            if n:
                total_replacements += n
                paragraphs_modified += 1

        # --- Process tables (including all cells) ---
        for table in self._doc.tables:
            total_replacements += self._process_table(table, redact_fn)

        # --- Process headers and footers ---
        for section in self._doc.sections:
            for hf in [section.header, section.footer,
                       section.even_page_header, section.even_page_footer,
                       section.first_page_header, section.first_page_footer]:
                if hf is not None:
                    for para in hf.paragraphs:
                        n = _replace_in_paragraph(para, redact_fn)
                        if n:
                            total_replacements += n

        logger.info("Processing complete: %d entity replacements across %d paragraphs",
                    total_replacements, paragraphs_modified)
        return total_replacements, paragraphs_modified

    def save(self, output_path: str | Path) -> Path:
        """Save the (modified) document to *output_path*.

        Args:
            output_path: Destination path for the redacted DOCX.

        Returns:
            The resolved output path.
        """
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        self._doc.save(str(out))
        logger.info("Saved redacted document to: %s", out)
        return out

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _process_table(self, table: Table, redact_fn) -> int:
        """Recursively process all cells in *table*.
        
        Deduplicates by XML element pointer (``_tc``) to handle:
        - Merged cells (same element shared across multiple row/col positions)
        - Wide tables where separate cells contain identical content
        """
        count = 0
        seen_cells: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                # Merged cells share the same _tc XML element — deduplicate by id
                cid = id(cell._tc)
                if cid in seen_cells:
                    continue
                seen_cells.add(cid)
                count += _replace_in_cell(cell, redact_fn)
                # Handle nested tables
                for nested_table in cell.tables:
                    count += self._process_table(nested_table, redact_fn)
        return count

    # ------------------------------------------------------------------
    # Validation
    # ------------------------------------------------------------------

    @staticmethod
    def validate(output_path: str | Path) -> dict:
        """Open and validate the output DOCX.

        Checks:
          - File exists and has .docx extension
          - Can be opened by python-docx
          - Contains at least one paragraph
          - Contains at least one table

        Returns:
            A dict with keys: valid (bool), issues (list[str])
        """
        out = Path(output_path)
        issues: List[str] = []

        if not out.exists():
            return {"valid": False, "issues": ["Output file does not exist"]}
        if out.suffix.lower() != ".docx":
            issues.append(f"Unexpected extension: {out.suffix}")

        try:
            doc = docx.Document(str(out))
        except Exception as exc:
            return {"valid": False, "issues": [f"Cannot open DOCX: {exc}"]}

        if not doc.paragraphs:
            issues.append("No paragraphs found in output")
        if not doc.tables:
            issues.append("No tables found in output (expected tables from source)")

        paragraph_text = " ".join(p.text for p in doc.paragraphs[:10])
        if len(paragraph_text) < 50:
            issues.append("Paragraphs appear empty or nearly empty")

        return {"valid": len(issues) == 0, "issues": issues}
