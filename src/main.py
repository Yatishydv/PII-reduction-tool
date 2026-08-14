"""
main.py — CLI entrypoint for the PII Redaction Tool.

Usage:
    python -m src.main --input "Red Herring Prospectus.docx" --output redacted_output.docx
    python -m src.main --evaluate --gold-annotations evaluation/gold_annotations.json \
                        --synthetic-cases evaluation/synthetic_cases.json
    python -m src.main --help

This module wires together:
    DocxProcessor → Redactor (all detectors) → ReplacementGenerator → output DOCX
    + Evaluator → metrics → evaluation_report.md
"""
from __future__ import annotations

import argparse
import json
import logging
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Logging setup — must run before any src imports
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("pii_redaction")


def _build_redactor(use_ner: bool = True):
    """Instantiate and return a fully-wired Redactor."""
    from src.detectors.regex_detectors import (
        EmailDetector, PhoneDetector, IPDetector,
        SSNDetector, CreditCardDetector, DOBDetector,
        PANDetector, AadhaarDetector, BankAccountDetector,
        IFSCDetector, DINDetector, GSTINDetector, CINDetector,
    )
    from src.detectors.address_detector import AddressDetector
    from src.detectors.entity_resolver import EntityResolver
    from src.redaction.replacement import ReplacementGenerator
    from src.redaction.redactor import Redactor

    detectors = [
        EmailDetector(),
        PhoneDetector(),
        IPDetector(),
        SSNDetector(),
        CreditCardDetector(),
        DOBDetector(),
        AddressDetector(),
        PANDetector(),
        AadhaarDetector(),
        BankAccountDetector(),
        IFSCDetector(),
        DINDetector(),
        GSTINDetector(),
        CINDetector(),
    ]

    if use_ner:
        try:
            from src.detectors.ner_detector import NERDetector
            ner = NERDetector()
            detectors.append(ner)
            logger.info("NER detector loaded successfully")
        except RuntimeError as exc:
            logger.warning("NER detector unavailable: %s — continuing without NER", exc)

    generator = ReplacementGenerator()
    resolver = EntityResolver()
    return Redactor(detectors=detectors, generator=generator, resolver=resolver)


def cmd_redact(args: argparse.Namespace) -> int:
    """Run the redaction pipeline on the input DOCX."""
    from src.docx_processor import DocxProcessor

    input_path = Path(args.input)
    output_path = Path(args.output)

    logger.info("=== PII Redaction Tool ===")
    logger.info("Input:  %s", input_path)
    logger.info("Output: %s", output_path)

    # Build redactor
    redactor = _build_redactor(use_ner=not args.no_ner)

    # Wrap redactor.redact as the redact_fn signature expected by DocxProcessor
    total_entities_detected: list = []

    def redact_fn(text: str):
        redacted, entities = redactor.redact(text)
        total_entities_detected.extend(entities)
        return redacted, entities

    # Load and process DOCX
    try:
        processor = DocxProcessor(input_path)
    except (FileNotFoundError, ValueError) as exc:
        logger.error("%s", exc)
        return 1

    replacements, paragraphs = processor.process(redact_fn)
    out_path = processor.save(output_path)

    # Validate output
    validation = DocxProcessor.validate(out_path)
    if validation["valid"]:
        logger.info("✅ Output DOCX validated successfully")
    else:
        logger.warning("⚠️  Output DOCX validation issues: %s", validation["issues"])

    # Summary stats
    from collections import Counter
    type_counts = Counter(e.label for e in total_entities_detected)
    logger.info("=== Redaction Summary ===")
    logger.info("Total entities detected & replaced: %d", len(total_entities_detected))
    for label, count in sorted(type_counts.items()):
        logger.info("  %-20s %d", label, count)
    logger.info("Output saved to: %s", out_path)

    # Post-redaction scan (residual PII check)
    # Note: synthetic replacements (example-corp.com, +91 900xxxxx) will also
    # trigger regex detectors — we filter those out to get true residuals.
    logger.info("=== Post-Redaction Scan ===")
    import docx as _docx
    from src.detectors.regex_detectors import (
        EmailDetector, PhoneDetector, IPDetector, SSNDetector, CreditCardDetector
    )
    from src.redaction.replacement import ReplacementGenerator

    # Get the set of all synthetic replacement values we generated
    replacement_values: set[str] = set()
    for v in redactor.get_replacement_map().values():
        replacement_values.add(v.strip().lower())

    redacted_doc = _docx.Document(str(out_path))
    residual_detectors = [EmailDetector(), PhoneDetector(), IPDetector(),
                          SSNDetector(), CreditCardDetector()]
    residual_count = 0
    residual_examples: list = []

    def is_synthetic(text: str) -> bool:
        """Return True if this detected value is one of our own synthetic replacements."""
        t = text.strip().lower()
        # Check against replacement map
        if t in replacement_values:
            return True
        # Also check our known synthetic domains
        if "example-corp.com" in t:
            return True
        # Synthetic phone prefix we generate: +91 90xxxxx
        import re as _re
        if _re.match(r'\+91\s+90\d', t):
            return True
        return False

    for para in redacted_doc.paragraphs:
        for det in residual_detectors:
            hits = det.detect(para.text)
            real_hits = [h for h in hits if not is_synthetic(h.text)]
            if real_hits:
                residual_count += len(real_hits)
                for h in real_hits[:2]:
                    residual_examples.append(f"{h.label}: {h.text[:50]}")
                logger.warning("Residual %s detected in output paragraph", det.label)

    for table in redacted_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for det in residual_detectors:
                    hits = det.detect(cell.text)
                    real_hits = [h for h in hits if not is_synthetic(h.text)]
                    if real_hits:
                        residual_count += len(real_hits)

    if residual_count == 0:
        logger.info("✅ No residual original PII detected in output")
    else:
        logger.warning("⚠️  %d residual original PII entities found in output", residual_count)
        for ex in residual_examples[:5]:
            logger.warning("  Example: %s", ex)

    # Save replacement map (local only, not committed to git)
    if args.save_map:
        map_path = output_path.parent / "replacement_map.json"
        full_map = redactor.get_replacement_map()
        serializable = {f"{k[0]}::{k[1]}": v for k, v in full_map.items()}
        with open(map_path, "w", encoding="utf-8") as f:
            json.dump(serializable, f, indent=2, ensure_ascii=False)
        logger.info("Replacement map saved to: %s (local only, not for distribution)", map_path)

    return 0


def cmd_evaluate(args: argparse.Namespace) -> int:
    """Run evaluation and generate report."""
    from src.evaluation.evaluator import Evaluator
    from src.evaluation.report import ReportGenerator
    import json

    logger.info("=== Evaluation Mode ===")

    redactor = _build_redactor(use_ner=not args.no_ner)
    gold_path = Path(args.gold_annotations)
    synthetic_path = Path(args.synthetic_cases)
    report_path = Path(args.report_output)

    evaluator = Evaluator(redactor, gold_path, synthetic_path)

    gold_results = evaluator.evaluate_gold()
    syn_results = evaluator.evaluate_synthetic()

    # Log results
    for label, m in sorted(gold_results.items()):
        logger.info("[GOLD] %-20s P=%.4f R=%.4f F1=%.4f TP=%d FP=%d FN=%d",
                    label, m.precision, m.recall, m.f1, m.tp, m.fp, m.fn)

    for label, m in sorted(syn_results.items()):
        logger.info("[SYNTHETIC] %-20s P=%.4f R=%.4f F1=%.4f TP=%d FP=%d FN=%d",
                    label, m.precision, m.recall, m.f1, m.tp, m.fp, m.fn)

    # Collect FP/FN examples
    import json
    with open(gold_path, "r", encoding="utf-8") as f:
        gold_data = json.load(f)
    with open(synthetic_path, "r", encoding="utf-8") as f:
        synthetic_data = json.load(f)

    gold_fps = evaluator.get_false_positives(gold_data)
    gold_fns = evaluator.get_false_negatives(gold_data)
    syn_fps = evaluator.get_false_positives(synthetic_data)
    syn_fns = evaluator.get_false_negatives(synthetic_data)

    all_fps = gold_fps + syn_fps
    all_fns = gold_fns + syn_fns

    # Save raw results JSON
    results_path = report_path.parent / "evaluation_results.json"
    with open(results_path, "w", encoding="utf-8") as f:
        json.dump(
            {
                "gold": {k: v.to_dict() for k, v in gold_results.items()},
                "synthetic": {k: v.to_dict() for k, v in syn_results.items()},
                "false_positives_count": len(all_fps),
                "false_negatives_count": len(all_fns),
            },
            f, indent=2
        )
    logger.info("Results saved to: %s", results_path)

    # Generate report
    reporter = ReportGenerator(
        gold_results=gold_results,
        synthetic_results=syn_results,
        gold_count=len(gold_data),
        synthetic_count=len(synthetic_data),
        false_positives=all_fps,
        false_negatives=all_fns,
    )
    reporter.write(report_path)
    logger.info("Evaluation report written to: %s", report_path)

    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="python -m src.main",
        description="PII Redaction Tool — Hybrid regex + NER approach",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Redact a DOCX file:
  python -m src.main --input "Red Herring Prospectus.docx" --output redacted_output.docx

  # Evaluate against annotation datasets:
  python -m src.main --evaluate \\
      --gold-annotations evaluation/gold_annotations.json \\
      --synthetic-cases evaluation/synthetic_cases.json

  # Redact without NER (faster, regex-only):
  python -m src.main --input file.docx --output out.docx --no-ner
        """,
    )

    # Mode
    mode_group = parser.add_mutually_exclusive_group()
    mode_group.add_argument(
        "--evaluate", action="store_true",
        help="Run evaluation mode (compute metrics from annotation files)",
    )

    # Redaction args
    parser.add_argument("--input", "-i", help="Input .docx file path")
    parser.add_argument("--output", "-o", default="redacted_output.docx",
                        help="Output .docx file path (default: redacted_output.docx)")

    # Evaluation args
    parser.add_argument(
        "--gold-annotations",
        default="evaluation/gold_annotations.json",
        help="Path to gold annotation JSON (default: evaluation/gold_annotations.json)",
    )
    parser.add_argument(
        "--synthetic-cases",
        default="evaluation/synthetic_cases.json",
        help="Path to synthetic test cases JSON (default: evaluation/synthetic_cases.json)",
    )
    parser.add_argument(
        "--report-output",
        default="evaluation/evaluation_report.md",
        help="Path to write evaluation report (default: evaluation/evaluation_report.md)",
    )

    # Options
    parser.add_argument(
        "--no-ner", action="store_true",
        help="Disable NER (use regex detectors only — faster but lower recall for names/orgs)",
    )
    parser.add_argument(
        "--save-map", action="store_true",
        help="Save replacement mapping to replacement_map.json (local only)",
    )
    parser.add_argument(
        "--verbose", "-v", action="store_true",
        help="Enable DEBUG logging",
    )

    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.evaluate:
        return cmd_evaluate(args)

    if not args.input:
        parser.error("--input is required for redaction mode (use --evaluate for evaluation)")

    return cmd_redact(args)


if __name__ == "__main__":
    sys.exit(main())
